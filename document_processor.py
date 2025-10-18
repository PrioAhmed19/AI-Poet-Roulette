import os
from typing import List
from pathlib import Path
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from config import CHUNK_SIZE, CHUNK_OVERLAP


class DocumentProcessor:
    """Process various document types and extract text content."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP
        )
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """
        Process a document and return chunked text.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of LangchainDocument objects
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text = self._extract_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self._extract_docx(file_path)
        elif file_ext == '.txt':
            text = self._extract_txt(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            text = self._extract_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Create document with metadata
        doc = LangchainDocument(
            page_content=text,
            metadata={
                "source": file_path,
                "file_type": file_ext
            }
        )
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        return chunks
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def _extract_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            return f"Error extracting text from image: {str(e)}"
